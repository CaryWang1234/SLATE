"""文件路由：多模态文件上传与解析。"""

from __future__ import annotations

import base64
import csv
import io
import json
import os
import re
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException, UploadFile
from fastapi.responses import FileResponse

router = APIRouter(prefix="/files", tags=["files"])

# 工具多模态输出（chart/qrcode 等）落盘目录，与 skills 保持一致
DATA_DIR = Path(os.environ.get("SLATE_DATA_DIR", Path(__file__).resolve().parent.parent.parent / "data"))
OUTPUTS_DIR = DATA_DIR / "outputs"

# 允许预览的输出文件类型 → MIME
OUTPUT_SERVE_TYPES = {
    ".svg": "image/svg+xml", ".png": "image/png", ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg", ".webp": "image/webp", ".gif": "image/gif",
    ".md": "text/markdown; charset=utf-8", ".json": "application/json; charset=utf-8",
    ".mp4": "video/mp4", ".webm": "video/webm",
}

# 支持的文本文件扩展名
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".html", ".htm", ".css",
    ".js", ".ts", ".jsx", ".tsx", ".py", ".java", ".c", ".cpp", ".h",
    ".json", ".yaml", ".yml", ".toml", ".xml", ".sql",
    ".sh", ".bat", ".ps1", ".rb", ".go", ".rs", ".php",
    ".csv", ".tsv", ".log", ".ini", ".cfg", ".conf",
}

# 支持的图片扩展名
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".svg"}

# 需后端库解析的文档类扩展名（docx/xlsx/pdf）
DOC_EXTENSIONS = {".docx"}
SHEET_EXTENSIONS = {".xlsx"}
PDF_EXTENSIONS = {".pdf"}

# 最大文件大小 (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024


@router.post("/upload")
async def upload_file(file: UploadFile) -> dict[str, Any]:
    """
    上传并解析文件。
    文本文件：提取内容文本。
    CSV：解析为结构化数据。
    图片：转为 base64 供多模态模型使用。
    """
    if not file.filename:
        return {"code": -1, "data": None, "message": "未提供文件名"}

    # 文件名消毒：去除路径分隔符和特殊字符
    filename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', file.filename)
    filename = filename.strip('. ')
    if not filename:
        filename = 'unnamed'
    ext = Path(filename).suffix.lower()
    content = await file.read()

    if len(content) > MAX_FILE_SIZE:
        return {"code": -1, "data": None, "message": f"文件过大（最大 {MAX_FILE_SIZE // 1024 // 1024}MB）"}

    # 文本文件
    if ext in TEXT_EXTENSIONS:
        return _parse_text_file(filename, ext, content)

    # 图片文件
    if ext in IMAGE_EXTENSIONS:
        return _parse_image_file(filename, ext, content)

    # 文档类：后端库解析（浏览器无法直接读取）
    if ext in DOC_EXTENSIONS:
        return _parse_docx(filename, content)
    if ext in SHEET_EXTENSIONS:
        return _parse_xlsx(filename, content)
    if ext in PDF_EXTENSIONS:
        return _parse_pdf(filename, content)

    # 其他文件：尝试作为文本读取
    try:
        text = content.decode("utf-8")
        return {
            "code": 0,
            "data": {
                "filename": filename,
                "type": "text",
                "content": text[:50000],
                "size": len(content),
                "truncated": len(content) > 50000,
            },
            "message": "ok",
        }
    except UnicodeDecodeError:
        return {"code": -1, "data": None, "message": f"不支持的文件类型: {ext}"}


def _parse_text_file(filename: str, ext: str, content: bytes) -> dict[str, Any]:
    """解析文本类文件。"""
    text = content.decode("utf-8", errors="replace")

    # CSV 特殊处理：解析为结构化数据
    if ext in (".csv", ".tsv"):
        return _parse_csv(filename, text, ext)

    return {
        "code": 0,
        "data": {
            "filename": filename,
            "type": "text",
            "content": text[:50000],
            "size": len(content),
            "line_count": text.count("\n") + 1,
            "truncated": len(text) > 50000,
        },
        "message": "ok",
    }


def _parse_csv(filename: str, text: str, ext: str) -> dict[str, Any]:
    """解析 CSV/TSV 文件。"""
    delimiter = "\t" if ext == ".tsv" else ","
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = list(reader)

    if not rows:
        return {
            "code": 0,
            "data": {"filename": filename, "type": "csv", "headers": [], "rows": [], "row_count": 0},
            "message": "ok",
        }

    headers = rows[0]
    data_rows = rows[1:101]  # 最多返回 100 行数据

    return {
        "code": 0,
        "data": {
            "filename": filename,
            "type": "csv",
            "headers": headers,
            "rows": data_rows,
            "row_count": len(rows) - 1,
            "truncated": len(rows) > 101,
        },
        "message": "ok",
    }


def _parse_image_file(filename: str, ext: str, content: bytes) -> dict[str, Any]:
    """解析图片文件，转为 base64。"""
    mime_map = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif",
        ".webp": "image/webp", ".bmp": "image/bmp",
        ".svg": "image/svg+xml",
    }
    mime_type = mime_map.get(ext, "image/png")
    b64 = base64.b64encode(content).decode("ascii")

    return {
        "code": 0,
        "data": {
            "filename": filename,
            "type": "image",
            "mime_type": mime_type,
            "base64": b64,
            "size": len(content),
        },
        "message": "ok",
    }


def _parse_docx(filename: str, content: bytes) -> dict[str, Any]:
    """解析 Word 文档：提取段落与表格文本。"""
    try:
        import docx
    except ImportError:
        return {"code": -1, "data": None, "message": "缺少 python-docx 依赖，无法解析 Word 文档"}

    doc = docx.Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            parts.append(" | ".join(cells))

    full = "\n".join(parts)
    return {
        "code": 0,
        "data": {
            "filename": filename,
            "type": "text",
            "content": full[:50000],
            "size": len(content),
            "truncated": len(full) > 50000,
        },
        "message": "ok",
    }


def _parse_xlsx(filename: str, content: bytes) -> dict[str, Any]:
    """解析 Excel 工作簿：前 5 个工作表、每表最多 200 行，转为制表符分隔文本。"""
    try:
        import openpyxl
    except ImportError:
        return {"code": -1, "data": None, "message": "缺少 openpyxl 依赖，无法解析 Excel 文件"}

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for ws in wb.worksheets[:5]:
            parts.append(f"[工作表: {ws.title}]")
            for row in ws.iter_rows(values_only=True, max_row=200):
                parts.append("\t".join("" if v is None else str(v) for v in row))
    finally:
        wb.close()

    full = "\n".join(parts)
    return {
        "code": 0,
        "data": {
            "filename": filename,
            "type": "text",
            "content": full[:50000],
            "size": len(content),
            "truncated": len(full) > 50000,
        },
        "message": "ok",
    }


def _parse_pdf(filename: str, content: bytes) -> dict[str, Any]:
    """解析 PDF：提取文本层（最多 50 页），扫描件无文本层时提示。"""
    try:
        import pdfplumber
    except ImportError:
        return {"code": -1, "data": None, "message": "缺少 pdfplumber 依赖，无法解析 PDF 文件"}

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages[:50]:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)

    full = "\n".join(parts)
    if not full.strip():
        return {
            "code": 0,
            "data": {
                "filename": filename,
                "type": "text",
                "content": "（该 PDF 未提取到文本层，可能为扫描件/纯图片 PDF）",
                "size": len(content),
                "truncated": False,
            },
            "message": "ok",
        }
    return {
        "code": 0,
        "data": {
            "filename": filename,
            "type": "text",
            "content": full[:50000],
            "size": len(content),
            "truncated": len(full) > 50000,
        },
        "message": "ok",
    }


@router.get("/output")
async def get_output_file(name: str) -> FileResponse:
    """提供 outputs 目录下工具产出图片的预览访问（仅限文件名，防路径穿越）。"""
    if not name or name != Path(name).name or "/" in name or "\\" in name or ".." in name:
        raise HTTPException(status_code=400, detail="非法文件名")
    ext = Path(name).suffix.lower()
    if ext not in OUTPUT_SERVE_TYPES:
        raise HTTPException(status_code=400, detail="不支持的文件类型")
    path = OUTPUTS_DIR / name
    if not path.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(str(path), media_type=OUTPUT_SERVE_TYPES[ext])
