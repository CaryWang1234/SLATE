"""MCP 工具：Word 撰写 —— 根据标题、章节结构与正文内容生成 .docx Word 文档。

参数:
- title      : 文档标题（默认"未命名文档"）
- author     : 作者署名（默认 SLATE）
- content    : 正文内容，支持简单标记语法：
               "# 文本"一级标题 / "## "二级 / "### "三级，
               "- 文本"无序列表 / "1. 文本"有序列表 / "> 文本"引用，
               其余为普通段落，空行分段
- sections   : 高级参数，JSON 数组精确控制章节，提供时忽略 content，格式：
               [{"heading": "章节标题", "level": 2, "paragraphs": ["段落..."], "bullets": ["列表项..."]}]
- output_dir : 输出目录（默认 SLATE 数据目录下的 outputs/）

实现：python-docx 生成文档，标题层级使用 Heading 1-3 样式，中文统一微软雅黑。
"""

from __future__ import annotations

import json
import os
import re
import time
from datetime import date
from pathlib import Path
from typing import Any

DATA_DIR = Path(os.environ.get("SLATE_DATA_DIR", Path(__file__).resolve().parent.parent.parent / "data"))

FONT_NAME = "微软雅黑"


def execute(
    title: str = "",
    author: str = "SLATE",
    content: str = "",
    sections: Any = None,
    output_dir: str = "",
    **_kw: Any,
) -> dict[str, Any]:
    """生成 .docx Word 文档，返回文件路径与统计信息。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"error": "缺少依赖 python-docx，请先执行: pip install python-docx"}

    doc_title = (title or "").strip() or "未命名文档"
    doc = Document()
    _setup_normal_style(doc)

    # ── 文档标题与署名 ──────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    _set_run_font(run)

    meta = " · ".join(x for x in [(author or "").strip(), date.today().strftime("%Y-%m-%d")] if x)
    if meta:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rm = pm.add_run(meta)
        rm.font.size = Pt(10)
        rm.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
        _set_run_font(rm)
        pm.paragraph_format.space_after = Pt(18)

    # ── 正文 ────────────────────────────────
    heading_count = 0
    paragraph_count = 0
    if sections:
        parsed, err = _parse_sections(sections)
        if err:
            return {"error": err}
        heading_count, paragraph_count = _render_sections(doc, parsed)
        if heading_count == 0 and paragraph_count == 0:
            return {"error": "sections 中没有可渲染的内容"}
    elif (content or "").strip():
        heading_count, paragraph_count = _render_content(doc, content)
    else:
        return {"error": "content 与 sections 至少提供一个"}

    # ── 保存 ────────────────────────────────
    out_dir = Path(output_dir.strip()) if (output_dir or "").strip() else DATA_DIR / "outputs"
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
        safe_name = re.sub(r'[\\/:*?"<>|\r\n]', "_", doc_title).strip()[:40] or "document"
        out_path = out_dir / f"{safe_name}_{time.strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(out_path))
    except OSError as e:
        return {"error": f"文件保存失败: {e}"}

    return {
        "file_path": str(out_path),
        "file_name": out_path.name,
        "size_bytes": out_path.stat().st_size,
        "heading_count": heading_count,
        "paragraph_count": paragraph_count,
        "note": "Word 文档已生成，可用 Word / WPS 打开。",
    }


# ── 渲染 ──────────────────────────────────────


def _render_content(doc: Any, content: str) -> tuple[int, int]:
    """解析带简单标记的正文文本。"""
    from docx.shared import Pt, RGBColor

    heading_count = 0
    paragraph_count = 0
    for raw in content.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if m:
            _add_heading(doc, m.group(2).strip(), len(m.group(1)))
            heading_count += 1
            continue

        stripped = line.lstrip()
        if stripped.startswith("- ") or stripped.startswith("* "):
            _add_bullet(doc, stripped[2:].strip())
            paragraph_count += 1
            continue
        if re.match(r"^\d+[.、)]\s+", stripped):
            _add_numbered(doc, re.sub(r"^\d+[.、)]\s+", "", stripped))
            paragraph_count += 1
            continue
        if stripped.startswith("> "):
            pq = doc.add_paragraph(style="Quote") if _has_style(doc, "Quote") else doc.add_paragraph()
            rq = pq.add_run(stripped[2:].strip())
            rq.italic = True
            rq.font.color.rgb = RGBColor(0x5C, 0x53, 0x47)
            rq.font.size = Pt(10.5)
            _set_run_font(rq)
            paragraph_count += 1
            continue

        _add_body(doc, line.strip())
        paragraph_count += 1

    return heading_count, paragraph_count


def _parse_sections(sections: Any) -> tuple[list | None, str]:
    """解析 sections 参数为列表，失败时返回错误信息。"""
    raw = sections
    if isinstance(raw, str):
        try:
            raw = json.loads(raw)
        except (ValueError, TypeError):
            return None, "sections 不是合法的 JSON 数组"
    if not isinstance(raw, list):
        return None, "sections 必须是 JSON 数组"
    return raw, ""


def _render_sections(doc: Any, items: list) -> tuple[int, int]:
    """渲染结构化章节列表。"""
    heading_count = 0
    paragraph_count = 0
    for item in items:
        if not isinstance(item, dict):
            continue
        heading = str(item.get("heading") or "").strip()
        if heading:
            level = item.get("level", 2)
            try:
                level = max(1, min(3, int(level)))
            except (TypeError, ValueError):
                level = 2
            _add_heading(doc, heading, level)
            heading_count += 1
        for text in item.get("paragraphs") or []:
            text = str(text).strip()
            if text:
                _add_body(doc, text)
                paragraph_count += 1
        for text in item.get("bullets") or []:
            text = str(text).strip()
            if text:
                _add_bullet(doc, text)
                paragraph_count += 1

    return heading_count, paragraph_count


# ── 排版辅助 ──────────────────────────────────


def _setup_normal_style(doc: Any) -> None:
    """正文默认样式：11pt 微软雅黑。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)


def _add_heading(doc: Any, text: str, level: int) -> None:
    """标题使用 Heading 样式但改为深灰近黑，避免默认蓝色。"""
    from docx.shared import RGBColor

    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        _set_run_font(run)


def _add_body(doc: Any, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run)
    p.paragraph_format.space_after = Pt(6)


def _add_bullet(doc: Any, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_font(run)


def _add_numbered(doc: Any, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    _set_run_font(run)


def _set_run_font(run: Any) -> None:
    """同时设置西文与中文（东亚）字体。"""
    from docx.oxml.ns import qn

    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)


def _has_style(doc: Any, name: str) -> bool:
    try:
        return any(s.name == name for s in doc.styles)
    except Exception:
        return False
